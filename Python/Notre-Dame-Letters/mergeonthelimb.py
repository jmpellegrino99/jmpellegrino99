from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import pandas as pd
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.enum.text import WD_BREAK

df = pd.read_csv('carve.csv')

df['ClassOf'] = df['ClassOf'].astype(int)
df['ClassOf'] = df['ClassOf'].astype(str)

df.fillna('', inplace=True)

def create_student_document(row):
    doc = Document()

    heading1 = doc.add_paragraph()
    r1 = heading1.add_run("Notre Dame High School Enrollment Contract: 2024-2025")
    r1.bold = True

    table = doc.add_table(rows=3,cols=2)
    table.style = 'Table Grid'

    cell_1 = table.cell(0,0)
    cell_2 = table.cell(1,0)
    cell_3 = table.cell(2,0)
    cell_4 = table.cell(0,1)
    cell_5 = table.cell(1,1)
    cell_6 = table.cell(2,1)

    cell_1.text = "Student Name:"
    cell_2.text = "Class of:"
    cell_3.text = "ID:"
    cell_4.text = f"{row['Last_Name']}, {row['First_Name']}"
    cell_5.text = f"{row['ClassOf']}"
    cell_6.text = f"{row['Student_Number']}"

    cell_1.paragraphs[0].runs[0].bold = True
    cell_2.paragraphs[0].runs[0].bold = True
    cell_3.paragraphs[0].runs[0].bold = True

    heading2 = doc.add_paragraph()
    r2 = heading2.add_run("\nAcceptance of 2024-2025 Enrollment Contract:")
    r2.bold = True

    # Add the paragraph to the document
    paragraph_text = "To remain enrolled at Notre Dame High School 2024 - 2025 academic year, the parents or legal guardians must submit this signed enrollment contract by "
    paragraph = doc.add_paragraph()

    # Add the first part of the paragraph
    run1 = paragraph.add_run(paragraph_text)

    # Add the bold part of the paragraph
    run2 = paragraph.add_run("March 15, 2024")
    run2.bold = True

    # Add the remaining part of the paragraph
    run3 = paragraph.add_run(". This must be accompanied by a non-refundable tuition deposit of $1,000. Upon submission of this contract and the required deposit, Notre Dame High School agrees to reserve a place for the academic year 2024-2025. This is a one-year contract. Students are issued a new enrollment contract each year.")

    if row['Amount1'] == "" and row['Amount2'] == "" and row['Amount3'] == "" and row['Amount4'] == "":
        table2 = doc.add_table(rows=3,cols=2)
        table2.style = 'Table Grid'

        cell2_1 = table2.cell(0,0)
        cell2_2 = table2.cell(1,0)
        cell2_3 = table2.cell(2,0)
        cell2_4 = table2.cell(0,1)
        cell2_5 = table2.cell(1,1)
        cell2_6 = table2.cell(2,1)

        cell2_1.text = "Tuition and Fees:"
        cell2_2.text = "Less Non-Refundable Enrollment Deposit:"
        cell2_3.text = "Net Payment Plan Amount:"
        cell2_4.text = f"{row['Tuition']}"
        cell2_5.text = "$1,000"
        cell2_6.text = f"{row['Net_Amount']}"

        cell2_1.paragraphs[0].runs[0].bold = True
        cell2_2.paragraphs[0].runs[0].bold = True
        cell2_3.paragraphs[0].runs[0].bold = True

    if row['Amount1'] != "" and row['Amount2'] == "" and row['Amount3'] == "" and row['Amount4'] == "":
        table2 = doc.add_table(rows=6,cols=2)
        table2.style = 'Table Grid'

        cell2_1 = table2.cell(0,0)
        cell2_2 = table2.cell(1,0)
        cell2_3 = table2.cell(2,0)
        cell2_4 = table2.cell(3,0)
        cell2_5 = table2.cell(4,0)
        cell2_6 = table2.cell(5,0)
        cell2_7 = table2.cell(0,1)
        cell2_8 = table2.cell(1,1)
        cell2_9 = table2.cell(2,1)
        cell2_10 = table2.cell(3,1)
        cell2_11 = table2.cell(4,1)
        cell2_12 = table2.cell(5,1)

        cell2_1.text = "Tuition and Fees:"
        cell2_2.text = ""
        cell2_3.text = f"{row['Award1']}"
        cell2_4.text = "Less Total Financial Aid:"
        cell2_5.text = "Less Non-Refundable Enrollment Deposit:"
        cell2_6.text = "Net Payment Plan Amount:"
        cell2_7.text = f"{row['Tuition']}"
        cell2_8.text = ""
        cell2_9.text = f"{row['Amount1']}"
        cell2_10.text = f"{row['Total_Aid']}"
        cell2_11.text = "$1,000"
        cell2_12.text = f"{row['Net_Amount']}"

        cell2_1.paragraphs[0].runs[0].bold = True
        cell2_3.paragraphs[0].runs[0].bold = True
        cell2_4.paragraphs[0].runs[0].bold = True
        cell2_5.paragraphs[0].runs[0].bold = True
        cell2_6.paragraphs[0].runs[0].bold = True

    if row['Amount1'] != "" and row['Amount2'] != "" and row['Amount3'] == "" and row['Amount4'] == "":
        table2 = doc.add_table(rows=7,cols=2)
        table2.style = 'Table Grid'

        cell2_1 = table2.cell(0,0)
        cell2_2 = table2.cell(1,0)
        cell2_3 = table2.cell(2,0)
        cell2_4 = table2.cell(3,0)
        cell2_5 = table2.cell(4,0)
        cell2_6 = table2.cell(5,0)
        cell2_7 = table2.cell(6,0)
        cell2_8 = table2.cell(0,1)
        cell2_9 = table2.cell(1,1)
        cell2_10 = table2.cell(2,1)
        cell2_11 = table2.cell(3,1)
        cell2_12 = table2.cell(4,1)
        cell2_13 = table2.cell(5,1)
        cell2_14 = table2.cell(6,1)

        cell2_1.text = "Tuition and Fees:"
        cell2_2.text = ""
        cell2_3.text = f"{row['Award1']}"
        cell2_4.text = f"{row['Award2']}"
        cell2_5.text = "Less Total Financial Aid:"
        cell2_6.text = "Less Non-Refundable Enrollment Deposit:"
        cell2_7.text = "Net Payment Plan Amount:"
        cell2_8.text = f"{row['Tuition']}"
        cell2_9.text = ""
        cell2_10.text = f"{row['Amount1']}"
        cell2_11.text = f"{row['Amount2']}"
        cell2_12.text = f"{row['Total_Aid']}"
        cell2_13.text = "$1,000"
        cell2_14.text = f"{row['Net_Amount']}"

        cell2_1.paragraphs[0].runs[0].bold = True
        cell2_3.paragraphs[0].runs[0].bold = True
        cell2_4.paragraphs[0].runs[0].bold = True
        cell2_5.paragraphs[0].runs[0].bold = True
        cell2_6.paragraphs[0].runs[0].bold = True
        cell2_7.paragraphs[0].runs[0].bold = True

    if row['Amount1'] != "" and row['Amount2'] != "" and row['Amount3'] != "" and row['Amount4'] == "":
        table2 = doc.add_table(rows=8,cols=2)
        table2.style = 'Table Grid'

        cell2_1 = table2.cell(0,0)
        cell2_2 = table2.cell(1,0)
        cell2_3 = table2.cell(2,0)
        cell2_4 = table2.cell(3,0)
        cell2_5 = table2.cell(4,0)
        cell2_6 = table2.cell(5,0)
        cell2_7 = table2.cell(6,0)
        cell2_8 = table2.cell(7,0)
        cell2_9 = table2.cell(0,1)
        cell2_10 = table2.cell(1,1)
        cell2_11 = table2.cell(2,1)
        cell2_12 = table2.cell(3,1)
        cell2_13 = table2.cell(4,1)
        cell2_14 = table2.cell(5,1)
        cell2_15 = table2.cell(6,1)
        cell2_16 = table2.cell(7,1)

        cell2_1.text = "Tuition and Fees:"
        cell2_2.text = ""
        cell2_3.text = f"{row['Award1']}"
        cell2_4.text = f"{row['Award2']}"
        cell2_5.text = f"{row['Award3']}"
        cell2_6.text = "Less Total Financial Aid:"
        cell2_7.text = "Less Non-Refundable Enrollment Deposit:"
        cell2_8.text = "Net Payment Plan Amount:"
        cell2_9.text = f"{row['Tuition']}"
        cell2_10.text = ""
        cell2_11.text = f"{row['Amount1']}"
        cell2_12.text = f"{row['Amount2']}"
        cell2_13.text = f"{row['Amount3']}"
        cell2_14.text = f"{row['Total_Aid']}"
        cell2_15.text = "$1,000"
        cell2_16.text = f"{row['Net_Amount']}"

        cell2_1.paragraphs[0].runs[0].bold = True
        cell2_3.paragraphs[0].runs[0].bold = True
        cell2_4.paragraphs[0].runs[0].bold = True
        cell2_5.paragraphs[0].runs[0].bold = True
        cell2_6.paragraphs[0].runs[0].bold = True
        cell2_7.paragraphs[0].runs[0].bold = True
        cell2_8.paragraphs[0].runs[0].bold = True

    if row['Amount1'] != "" and row['Amount2'] != "" and row['Amount3'] != "" and row['Amount4'] != "":
        table2 = doc.add_table(rows=9,cols=2)
        table2.style = 'Table Grid'

        cell2_1 = table2.cell(0,0)
        cell2_2 = table2.cell(1,0)
        cell2_3 = table2.cell(2,0)
        cell2_4 = table2.cell(3,0)
        cell2_5 = table2.cell(4,0)
        cell2_6 = table2.cell(5,0)
        cell2_7 = table2.cell(6,0)
        cell2_8 = table2.cell(7,0)
        cell2_9 = table2.cell(8,0)
        cell2_10 = table2.cell(0,1)
        cell2_11 = table2.cell(1,1)
        cell2_12 = table2.cell(2,1)
        cell2_13 = table2.cell(3,1)
        cell2_14 = table2.cell(4,1)
        cell2_15 = table2.cell(5,1)
        cell2_16 = table2.cell(6,1)
        cell2_17 = table2.cell(7,1)
        cell2_18 = table2.cell(8,1)

        cell2_1.text = "Tuition and Fees:"
        cell2_2.text = ""
        cell2_3.text = f"{row['Award1']}"
        cell2_4.text = f"{row['Award2']}"
        cell2_5.text = f"{row['Award3']}"
        cell2_6.text = f"{row['Award4']}"
        cell2_7.text = "Less Total Financial Aid:"
        cell2_8.text = "Less Non-Refundable Enrollment Deposit:"
        cell2_9.text = "Net Payment Plan Amount:"
        cell2_10.text = f"{row['Tuition']}"
        cell2_11.text = ""
        cell2_12.text = f"{row['Amount1']}"
        cell2_13.text = f"{row['Amount2']}"
        cell2_14.text = f"{row['Amount3']}"
        cell2_15.text = f"{row['Amount4']}"
        cell2_16.text = f"{row['Total_Aid']}"
        cell2_17.text = "$1,000"
        cell2_18.text = f"{row['Net_Amount']}"

        cell2_1.paragraphs[0].runs[0].bold = True
        cell2_3.paragraphs[0].runs[0].bold = True
        cell2_4.paragraphs[0].runs[0].bold = True
        cell2_5.paragraphs[0].runs[0].bold = True
        cell2_6.paragraphs[0].runs[0].bold = True
        cell2_7.paragraphs[0].runs[0].bold = True
        cell2_8.paragraphs[0].runs[0].bold = True
        cell2_9.paragraphs[0].runs[0].bold = True

    paragraph2 = doc.add_paragraph()

    run4 = paragraph2.add_run("\nPayment Plans: ")
    run4.bold = True

    run5 = paragraph2.add_run("Families may elect one of three payment plans. Please select one of the following options.\n\n")

    run6 = paragraph2.add_run("____ Plan 1 - One annual payment for the total tuition and fees due on May 1, 2024.\n\n")

    run7 = paragraph2.add_run("____ Plan 2 - Two semi-annual installments with payments due on May 1, 2024 and November 1, 2024.\n\n")

    run8 = paragraph2.add_run("____ Plan 3 - Ten monthly installments with payments due monthly from May 1, 2024 through February 1, 2025.\n\n")

    run9 = paragraph2.add_run("Account Administration: ")
    run9.bold = True

    run10 = paragraph2.add_run("Your tuition account and incidental billing will be administered by Blackbaud Tuition Management. Below please provide an email address for each responsible parent or guardian.\n\nEMail (Parent/Guardian A): ___________________________ EMail (Parent/Guardian B): ___________________________\n\nParent/Guardian A Name: _______________________________________________________________________\n\nParent/Guardian A Signature: _______________________________ Date: _______________________________\n\nParent/Guardian B Name: _______________________________________________________________________\n\nParent/Guardian B Signature: _______________________________ Date: _______________________________\n\nNotre Dame of West Haven: _______________________________________________________________________\n\n")

    run11 = paragraph2.add_run("Please review the information on the reverse side before signing and returning.")
    run11.bold = True

    paragraph2.add_run().add_break(break_type=WD_BREAK.PAGE)

    run12 = paragraph2.add_run("Late Payment Policy:\n")
    run12.bold = True

    run13 = paragraph2.add_run("Tuition Management will impose a $40 late fee on tuition payments not made by the due date.\n\n")

    run14 = paragraph2.add_run("If tuition obligations are not current, Notre Dame may prohibit the student from participating in academic and/or co-curricular activities, and grades and transcripts will be withheld. If a balance remains unpaid for more than thirty (30) days, Notre Dame may dismiss the student. Please note that a dismissal does not relieve the undersigned of the obligation to pay the full tuition due for the entire 2024-2025 academic year. The undersigned parents/guardians are liable for any costs incurred by Notre Dame High School in the enforcement of this enrollment contract.\n\n")

    run15 = paragraph2.add_run("Cancellation:\n")
    run15.bold = True

    run16 = paragraph2.add_run("Enrollment may be canceled by the undersigned in writing without penalty (beyond the non-refundable tuition deposit) prior to June 30, 2024. If the student withdraws after June 20, 2024, the undersigned is obligated to pay 65 percent of the full year’s tuition. If the student withdraws after October 31, 2024, the undersigned is responsible for the full amount of tuition and fees. Notre Dame High School will not provide transcripts or records for a student withdrawing from Notre Dame until all payments or amounts due are paid in full.\n\n")

    run17 = paragraph2.add_run("Rules and Regulations:\n")
    run17.bold = True

    run18 = paragraph2.add_run("The student shall uphold and abide by the rules and regulations of Notre Dame High School. Notre Dame reserves the right to nullify the enrollment contract prior to the start of the academic year, dismiss, or suspend the student at any time if such action is deemed in the best interest of Notre Dame High School. In addition, the parent/guardian may be asked to withdraw the student from the school if they are not supportive of Notre Dame’s goals and compliant with its policies. Such actions shall not release the undersigned of their financial obligations in this contract.\n\n")

    run19 = paragraph2.add_run("General Release:\n")
    run19.bold = True

    run20 = paragraph2.add_run("The undersigned recognize that the student may get injured during athletic or other school activities. The undersigned release and hold harmless Notre Dame High School, its agents, and employees from all claims, damages, and other liability for injury to the student where such claims, damages or other liability are not the result of gross negligence by Notre Dame High School, its agents, or employees.\n\n")

    run21 = paragraph2.add_run("Entire Agreement:\n")
    run21.bold = True

    run22 = paragraph2.add_run("The terms of this enrollment contract and the Student Handbook, and all forms attached thereto constitute the full and complete agreement between the parties and are the conditions upon which the student has been enrolled. No other verbal or written agreement shall, in any way, vary or alter any provision of this agreement unless both parties consent to vary or alter any provision of this contract in writing. Any prior oral or written agreements between the parties are merged into this contract and extinguished. No course of dealing between the parties shall in any way vary or alter the terms and conditions of this contract.\n\n")

    run23 = paragraph2.add_run("Severability:\n")
    run23.bold = True

    run24 = paragraph2.add_run("If any provision of this contract is deemed invalid, illegal, or unenforceable, all other conditions and provisions of this contract shall remain in full force and effect.\n\n")

    run25 = paragraph2.add_run("Applicable Law:\n")
    run25.bold = True

    run26 = paragraph2.add_run("This enrollment contract shall be interpreted and enforced in accordance with the laws of the State of Connecticut.\n\nBy signing this document, the individual acknowledges that he/she has read, understands, and hereby agrees to all of the terms and conditions of this enrollment contract and to be financially liable to pay the tuition, fees, and charges described in this enrollment contract.")

    for paragraph in doc.paragraphs:
        paragraph.style.font.name = 'Times New Roman'
        paragraph.style.font.size = Pt(10)

    sections = doc.sections
    for section in sections:
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(1.5)
        section.bottom_margin = Inches(0.5)

    para = doc.paragraphs[0]
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.save(f"LimbLetters/{row['Last_Name']}_{row['First_Name']}_{row['Student_Number']}_{row['Index']}.docx")

for index, row in df.iterrows():
    create_student_document(row)
